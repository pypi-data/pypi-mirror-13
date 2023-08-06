
# 1.0.0     First version
# 1.0.4     Fixed "unkown character" issues
# 1.0.5     Allows middle initial
# 1.0.6     Added ability to get titles, added some timeouts b/c WL is slower now, added time out messages
# 1.0.7     Fixed references timeout error. Some documents don't have a references tab. Just set it to zero and move on.
#   Need to add ability to enter multiple authors, separated by newlines, so can just do them all at once, duh

import os
import getpass
import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.common.exceptions import TimeoutException, NoSuchElementException, ElementNotVisibleException,\
StaleElementReferenceException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support import expected_conditions as EC

VERSION_NUM = "1.0.7"

if os.name == 'nt':
    DOWNLOAD_DIR = "C:"+os.environ['HOMEPATH']+"\\Downloads"
else:
    DOWNLOAD_DIR = os.environ['HOME']+"/Downloads"

def get_westlaw_bibliography(driver, author_name, get_titles):
    # tell the user what's happening
    print("Searching WestLaw for the bibliography...")

    # initialize by splitting the author_name
    author_elements = author_name.split(" ")
    number_of_entries = []

    # go to westlaw
    driver.get("http://a.next.westlaw.com")

    # determine whether logged in or not. If not, log in.
    if "Signon" in driver.title:
        print("getbib has not logged into WestLaw yet...")
        username = input("Please enter your WestLaw username: ")
        password = getpass.getpass("Password: ")
        elem = driver.find_element_by_name("Username")
        elem.send_keys(username)
        elem = driver.find_element_by_name("Password")
        elem.send_keys(password)
        elem.send_keys(Keys.RETURN)
        try:
            # press the continue button
            elem = driver.find_element_by_id("co_clientIDContinueButton")
            elem.click()
        except NoSuchElementException:
            print("Unable to login.\nAre you sure your username and password are correct")
            return

    print("Locating...")

    # wait for the "Secondary Sources" box to appear
    try:
        wait = WebDriverWait(driver, 8)
        elem = wait.until(EC.presence_of_element_located((By.PARTIAL_LINK_TEXT, "Secondary Sources")))
        elem.click()
    except TimeoutException:
        print("Timed out waiting for secondary sources to appear. Please check connection and try again.")
        return
    except NoSuchElementException:
        print("Cannot find secondary sources link. Please check connection and try again.")
        return

    # wait for "Law Reviews & Journals" to appear
    try:
        wait = WebDriverWait(driver, 8)
        elem = wait.until(EC.presence_of_element_located((By.PARTIAL_LINK_TEXT, "Law Reviews &")))
        elem.click()
    except TimeoutException:
        print("Timed out waiting for link to appear. Please check connection and try again.")
        return
    except NoSuchElementException:
        print("Cannot find secondary sources link. Please check connection and try again.")
        return

    # look for the search box, input search and click button
    try:
        wait = WebDriverWait(driver, 8)
        elem = wait.until(EC.presence_of_element_located((By.ID, "searchInputId")))
        if len(author_elements) == 2:
            elem.send_keys("adv: AU("+author_elements[0]+" /2 "+author_elements[1]+")")
        elif len(author_elements) == 3:
            elem.send_keys("adv: AU("+author_elements[0]+" /2 "+author_elements[1]+" /2 "+author_elements[2]+")")
    except TimeoutException:
        print("Timed out looking for search box. Please check connection and try again.")
        return
    except NoSuchElementException:
        print("Cannot find search text box. Please check connection and try again.")
        return

    # look for the search button
    try:
        wait = WebDriverWait(driver, 8)
        elem = wait.until(EC.presence_of_element_located((By.ID, "searchButton")))
        elem.click()
    except TimeoutException:
        print("Timed out looking for search button. Please check connection and try again.")
        return
    except NoSuchElementException:
        print("Cannot find search button. Please check connection and try again.")
        return

    # organize by date
    # first find sort option
    try:
        wait = WebDriverWait(driver, 8)
        elem = wait.until(EC.presence_of_element_located((By.ID, "co_search_sortOptions")))
        # then click date
        try:
            for option in elem.find_elements_by_tag_name("option"):
                if "Date" in option.text:
                    option.click()                
        except NoSuchElementException:
            print("Cannot sort by date.")
            return    
        except TimeoutException:
            print("Timed out trying to find sort options. Please check connection and try again.")
            return
        except StaleElementReferenceException:
            print("Already sorted by date...")
    except NoSuchElementException:
        print("Please try another search.")
        return
    except TimeoutException:
        print("Timed out while trying to organize by date. Check connection, or try another search.")

    # Get the number of law review articles
    try:
        wait = WebDriverWait(driver, 8)
        elem = wait.until(EC.presence_of_element_located((By.ID, "co_facetLaw Reviews & Journals")))
        number_of_entries = elem.text.split("\n")
        print("Number of entries: %s" % number_of_entries[0])
    except TimeoutException:
        print("Timed out while waiting to get number of entries. Please check connection and try again.")
        return
    except NoSuchElementException:
        print("Cannot find number of entries. Please check connection and try again.")
        return
    except StaleElementReferenceException:
        print("Make sure number of entries is correct...")
        pass

    # ask the user how many entries to grab
    if int(number_of_entries[0]) < 21:
        max_number_of_entries = int(number_of_entries[0])
    elif number_of_entries[0] == '':
        print("Could not find any entries. May have timed out. Please try again.")
        return
    else:
        max_number_of_entries = int(get_max_number_of_entries())

    # create necessary variables
    bibliography_list = []

    # create a word document in memory
    document = Document()

    # access the first (only) section
    sections = document.sections
    section = sections[0]

    # set the margins
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # create the first paragraph, access the run, set font type and size
    run = document.add_paragraph().add_run()
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # click the first entry
    try:
        time.sleep(2)
        wait = WebDriverWait(driver, 10)
        elem = wait.until(EC.presence_of_element_located((By.ID, "cobalt_result_analytical_title1")))
        elem.click()
    except TimeoutException:
        print("Timed out attempting to click article title. Please check connection and try again.")
        return
    except NoSuchElementException:
        print("Cannot find search button. Please check connection and try again.")
        return

    # For the number of law review articles, do the process
    # would be better to scroll through after first search clicked
    for x in range(1, max_number_of_entries + 1):
        # get the review name
        try:
            wait = WebDriverWait(driver, 8)
            elem = wait.until(EC.presence_of_element_located((By.CLASS_NAME, "co_headtext")))
            current_journal = elem.text
            print 
        except TimeoutException:
            print("Timed out finding name of journal. Please check connection and try again.")
            return
        except NoSuchElementException:
            print("Cannot find title. Please check connection and try again.")
            return

        if get_titles == True:
            # get the title
            try:
                wait = WebDriverWait(driver, 8)
                elem = wait.until(EC.presence_of_element_located((By.ID, "title")))
                current_title = elem.text
                print 
            except TimeoutException:
                print("Timed out trying to find title. Please check connection and try again.")
                return
            except NoSuchElementException:
                print("Cannot find title. Please check connection and try again.")
                return

        # get the number of references
        try:
            # not always getting it, so sleep for one-half second, adjust as need be
            time.sleep(.5)
            wait = WebDriverWait(driver, 8)
            elem = wait.until(EC.presence_of_element_located((By.ID, "coid_relatedInfo_kcCitingReferences_link")))
            # first splt should be "Citing References (" and "#)"
            first_split = elem.text.split("(")
            # second split should only contain the number, splitting first_split[1] into two parts
            second_split = first_split[1].split(")")
            current_refs = second_split[0]
            print 
        except TimeoutException:
            print("Timed out trying to find number of references. Setting references to 0.")
            current_refs = "0"
        except NoSuchElementException:
            print("Cannot find number of references. Please check connection and try again.")
            return

        # get the date
        try:
            wait = WebDriverWait(driver, 8)
            elem = wait.until(EC.presence_of_element_located((By.CLASS_NAME, "co_date")))
            current_date = elem.text
            print 
        except TimeoutException:
            print("Timed out trying to find date. Please check connection and try again.")
            return
        except NoSuchElementException:
            print("Cannot find date. Please check and try again.")
            return

        if get_titles == True:
            print("%i.: %s, %s, %s, %s %s" % (x, current_journal, current_title, current_date, current_refs, "reference" if int(current_refs) == 1 else "references"))
            run.add_text(str(x)+". "+current_journal+", "+current_title+", "+current_date+", "+current_refs+" "+("reference" if int(current_refs) == 1 else "references"))
            run.add_break()
            # update the dictionary
            bibliography_list.append( (current_journal, current_title, current_date, current_refs) )
        else:
            print("%i.: %s, %s, %s %s" % (x, current_journal, current_date, current_refs, "reference" if int(current_refs) == 1 else "references"))
            run.add_text(str(x)+". "+current_journal+", "+current_date+", "+current_refs+" "+("reference" if int(current_refs) == 1 else "references"))
            run.add_break()
            # update the dictionary
            bibliography_list.append( (current_journal, current_date, current_refs) )

        # go to the next search page
        try:
            # find the "next" button
            wait = WebDriverWait(driver, 8)
            elem = wait.until(EC.presence_of_element_located((By.ID, "co_documentFooterResultsNavigationNext")))
            elem.click()
        except TimeoutException:
            print("Timed out trying to go to next article. Please check connection and try again.")
            return
        except NoSuchElementException:
            print("Cannot find search button. Please check connection and try again.")
            return

    print("Writing output to document...")
    os.chdir(DOWNLOAD_DIR)
    document.save(author_name+" Bibliography.docx")

    return


def get_driver(profile):
    # get firefox driver (start firefox)
    # using profile created in create_browser_profile()

    driver = webdriver.Firefox(firefox_profile=profile)
    return driver
 
def introduction():

    # clear the terminal screen
    os.system('cls' if os.name == 'nt' else 'clear')

    # print getbib introductio
    print("getbib "+VERSION_NUM+" (c) 2015")

    # give an appropriate greeting, dependent on the time of day
    current_time = datetime.now().time()
    if current_time.hour < 12:
        print("Good morning, Reviewer!")
    elif 12 <= current_time.hour < 17:
        print("Good afternoon, Reviewer!")
    else:
        print("Working the night shift, eh?")

    return
    
    
def create_browser_profile():
    # change profile to remove download dialog box
    profile = webdriver.FirefoxProfile()
    profile.set_preference('browser.download.folderList', 2) # custom location
    profile.set_preference('browser.download.manager.useWindow', False)
    profile.set_preference('browser.download.manager.showWhenStarting', False)
    profile.set_preference('browser.download.dir', DOWNLOAD_DIR)
    profile.set_preference('browser.download.manager.focusWhenStarting', False)
    profile.set_preference('browser.helperApps.alwaysAsk.force', False)
    profile.set_preference('services.sync.prefs.sync.browser.download.manager.showWhenStarting', False)
    profile.set_preference('pdfjs.disabled', True)
    profile.set_preference('browser.helperApps.neverAsk.saveToDisk', 'application/pdf')
    return profile

def get_max_number_of_entries():

    while True:
        max_number_of_entries = input("Please enter the maximum number of entries you would like document: ")
           
        # make sure only the proper characters
        if not re.match('^[0-9]+$', max_number_of_entries):
            print("Please enter a number.")
        else:
            return max_number_of_entries

def main():
    introduction()
    profile = create_browser_profile()
    driver = get_driver(profile)

    while True:
        author_name = input("Please enter the author's first and last name: ")

        # make sure only the proper characters
        if not re.match('^[\'A-Za-z .-]+$', author_name):
            print("Unknown characters. Please try again.")
        elif author_name.lower() == "quit":
            break
        else:
            get_titles = input("Do you want to retrieve titles? (y/n): ")
            while True:
                if not re.match('^[ynYn]+$', get_titles):
                    get_titles = input("Invalid answer. (y/n): ")
                # if not, just go back
                if get_titles.lower() == 'n':
                    get_titles = False
                    break
                if get_titles.lower() == 'y':
                    get_titles = True
                    break
            get_westlaw_bibliography(driver, author_name, get_titles)

    driver.close()
    return

main()



